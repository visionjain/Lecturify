import io
import os
import streamlit as st
from gtts import gTTS
from io import BytesIO
from audio_recorder_streamlit import audio_recorder
import speech_recognition as sr
import google.generativeai as genai
import markdown
from docx import Document
from io import BytesIO
from dotenv import load_dotenv

load_dotenv()
# Hide Streamlit footer
hide_streamlit_style = """
<style>
    footer {visibility: hidden;}
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# Configure Google Generative AI
genai.configure(api_key=os.getenv("GENAI_API_KEY"))

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

# Set up the generative model
model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config,
    system_instruction="Professional and concise",
)

chat_session = model.start_chat(
    history=[{"role": "user", "parts": [""]}]
)

def main():
    st.title("Lecturify - Lecture to Notes/Quiz Generator")
    st.header("Record and Transcribe Lecture")

    # Step 1: Record Audio
    audio_bytes = audio_recorder(pause_threshold=40)
    if audio_bytes:
        if len(audio_bytes) > 8000:
            st.success("Audio captured correctly")
        else:
            st.warning("Audio captured incorrectly, please try again.")
        st.audio(audio_bytes, format="audio/wav")
        st.session_state.audio_bytes = audio_bytes

    # Step 2: Transcribe Audio
    if "transcript" not in st.session_state:
        st.session_state.transcript = ""

    if st.button("Transcribe Audio") and 'audio_bytes' in st.session_state and len(st.session_state.audio_bytes) > 0:
        audio_file = io.BytesIO(st.session_state.audio_bytes)
        audio_file.name = "temp_audio_file.wav"

        recognizer = sr.Recognizer()
        with sr.AudioFile(audio_file) as source:
            audio_data = recognizer.record(source)
            try:
                st.session_state.transcript = recognizer.recognize_google(audio_data)
            except sr.UnknownValueError:
                st.warning("Could not understand the audio.")
            except sr.RequestError as e:
                st.warning(f"Could not request results from Google Speech Recognition service; {e}")

    # Display the transcription text area with editable content
    st.markdown("***Transcription:***")
    final_transcript = st.text_area(
        "Final Lecture Transcript", 
        st.session_state.transcript, 
        height=200, 
        key="transcript_display"
    )

    # Step 3: Use Transcript to Generate Lecture Material
    if final_transcript:
        st.header("Generate Lecture Material")

        def pdf_response(lecture_transcript):
            response = chat_session.send_message(f"Generate lecture notes. Use markdown format: {lecture_transcript}")
            return response.text, markdown.markdown(response.text)

        def cheatsheet_response(lecture_transcript):
            response = chat_session.send_message(f"Generate a cheatsheet with at least 30 bullet points from the lecture: {lecture_transcript}")
            return response.text, markdown.markdown(response.text)

        def flashcard_response(lecture_transcript):
            response = chat_session.send_message(f"Generate flashcards in the form of single line pointers: {lecture_transcript}")
            return response.text

        def quiz_response(lecture_transcript):
            response = chat_session.send_message(f"Generate 7-8 multiple choice questions with correct and incorrect answers: {lecture_transcript}")
            return response.text

        # Generate Notes
        if st.button("Generate Lecture Notes"):
            raw_response, formatted_notes = pdf_response(final_transcript)
            st.markdown("***Generated Lecture Notes:***")
            st.write(raw_response)

            # Create Word Document for Lecture Notes
            doc = Document()
            doc.add_heading('Lecture Notes', 0)
            doc.add_paragraph(raw_response)

            # Save Word document to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button("Download Lecture Notes as Word", buffer, file_name="lecture_notes.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Generate Cheatsheet
        if st.button("Generate Cheatsheet"):
            raw_response, formatted_cheatsheet = cheatsheet_response(final_transcript)
            st.markdown("***Generated Cheatsheet:***")
            st.write(raw_response)

            # Create Word Document for Cheatsheet
            doc = Document()
            doc.add_heading('Cheatsheet', 0)
            doc.add_paragraph(raw_response)

            # Save Word document to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button("Download Cheatsheet as Word", buffer, file_name="cheatsheet.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Generate Flashcards
        if st.button("Generate Pointers"):
            flashcards = flashcard_response(final_transcript)
            st.markdown("***Generated Flashcards:***")
            st.write(flashcards)

            # Create Word Document for Flashcards
            doc = Document()
            doc.add_heading('Flashcards', 0)
            doc.add_paragraph(flashcards)

            # Save Word document to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button("Download Pointers as Word", buffer, file_name="flashcards.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Generate Quiz
        if st.button("Generate Quiz"):
            quiz = quiz_response(final_transcript)
            st.markdown("***Generated Quiz:***")
            st.write(quiz)

            # Create Word Document for Quiz
            doc = Document()
            doc.add_heading('Quiz', 0)
            doc.add_paragraph(quiz)

            # Save Word document to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button("Download Quiz as Word", buffer, file_name="quiz.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == '__main__':
    main()
